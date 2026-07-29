import io
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, Any

def generate_builder_resume_docx(data: Dict[str, Any], template_style: str = "Classic Executive") -> bytes:
    """
    Generates a native Microsoft Word (.docx) document formatted according to the selected template style.
    
    Templates:
    - Classic Executive: Traditional layout, Dark Blue headings, Times New Roman
    - Modern Tech: Modern layout, Emerald Green headings, Calibri
    - Minimalist Clean: Sleek layout, Dark Slate headings, Arial
    - Creative Professional: Dynamic layout, Indigo Purple headings, Segoe UI / Arial
    """
    doc = Document()
    
    # Define Template Styling Properties
    if template_style == "Modern Tech":
        font_name = 'Calibri'
        primary_color = RGBColor(16, 185, 129)  # Emerald Green
        secondary_color = RGBColor(15, 23, 42)  # Dark Slate
    elif template_style == "Minimalist Clean":
        font_name = 'Arial'
        primary_color = RGBColor(51, 65, 85)   # Slate Blue
        secondary_color = RGBColor(15, 23, 42)  # Dark
    elif template_style == "Creative Professional":
        font_name = 'Arial'
        primary_color = RGBColor(99, 102, 241)  # Indigo
        secondary_color = RGBColor(30, 41, 59)  # Slate
    else:
        # Default: Classic Executive
        font_name = 'Times New Roman'
        primary_color = RGBColor(30, 58, 138)  # Classic Deep Blue
        secondary_color = RGBColor(17, 24, 39)  # Dark Gray

    # Page Margins (0.75 inch)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # 1. Candidate Name (Header)
    name = data.get("Name", "Candidate Name")
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = name_p.add_run(name)
    run_name.font.name = font_name
    run_name.font.size = Pt(24)
    run_name.font.bold = True
    run_name.font.color.rgb = primary_color

    # 2. Contact Information Subtitle
    email = data.get("Email", "")
    phone = data.get("Phone", "")
    location = data.get("Location", "")
    linkedin = data.get("LinkedIn", "")

    contact_bits = [b for b in [email, phone, location, linkedin] if b]
    if contact_bits:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_contact = contact_p.add_run(" | ".join(contact_bits))
        run_contact.font.name = font_name
        run_contact.font.size = Pt(9.5)
        run_contact.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_section_header(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title.upper())
        run.font.name = font_name
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = primary_color

    # 3. Professional Summary
    summary = data.get("Professional_Summary", "")
    if summary:
        add_section_header("Professional Summary")
        p_sum = doc.add_paragraph()
        p_sum.paragraph_format.space_after = Pt(8)
        run_sum = p_sum.add_run(summary)
        run_sum.font.name = font_name
        run_sum.font.size = Pt(10)
        run_sum.font.color.rgb = secondary_color

    # 4. Technical & Core Skills
    skills = data.get("Skills", [])
    if skills:
        add_section_header("Technical & Core Competencies")
        p_skills = doc.add_paragraph()
        p_skills.paragraph_format.space_after = Pt(8)
        skills_str = " • ".join(skills) if isinstance(skills, list) else str(skills)
        run_sk = p_skills.add_run(skills_str)
        run_sk.font.name = font_name
        run_sk.font.size = Pt(10)
        run_sk.font.color.rgb = secondary_color

    # 5. Professional Experience
    exp_list = data.get("Experience", [])
    if exp_list:
        add_section_header("Professional Experience")
        for exp in exp_list:
            if isinstance(exp, dict):
                comp = exp.get("Company", "Company")
                role = exp.get("Role", "Role")
                dur = exp.get("Duration", "")

                p_exp = doc.add_paragraph()
                p_exp.paragraph_format.space_before = Pt(4)
                p_exp.paragraph_format.space_after = Pt(2)
                
                run_role = p_exp.add_run(f"{role}")
                run_role.font.name = font_name
                run_role.font.size = Pt(10.5)
                run_role.font.bold = True
                run_role.font.color.rgb = secondary_color

                run_comp = p_exp.add_run(f" — {comp}")
                run_comp.font.name = font_name
                run_comp.font.size = Pt(10)
                run_comp.font.italic = True
                run_comp.font.color.rgb = secondary_color

                if dur:
                    run_dur = p_exp.add_run(f" ({dur})")
                    run_dur.font.name = font_name
                    run_dur.font.size = Pt(9.5)
                    run_dur.font.color.rgb = RGBColor(100, 116, 139)

                bullets = exp.get("Bullet_Points") or exp.get("Description") or []
                if isinstance(bullets, list):
                    for b in bullets:
                        bp = doc.add_paragraph(style='List Bullet')
                        bp.paragraph_format.space_after = Pt(2)
                        run_b = bp.add_run(str(b))
                        run_b.font.name = font_name
                        run_b.font.size = Pt(9.5)
                        run_b.font.color.rgb = secondary_color
                elif bullets:
                    bp = doc.add_paragraph(style='List Bullet')
                    bp.paragraph_format.space_after = Pt(2)
                    run_b = bp.add_run(str(bullets))
                    run_b.font.name = font_name
                    run_b.font.size = Pt(9.5)
                    run_b.font.color.rgb = secondary_color
            else:
                p_exp = doc.add_paragraph(style='List Bullet')
                run_e = p_exp.add_run(str(exp))
                run_e.font.name = font_name
                run_e.font.size = Pt(9.5)

    # 6. Projects
    proj_list = data.get("Projects", [])
    if proj_list:
        add_section_header("Key Technical Projects")
        for proj in proj_list:
            if isinstance(proj, dict):
                title = proj.get("Title", "Project Title")
                tech = proj.get("TechStack", "")

                p_proj = doc.add_paragraph()
                p_proj.paragraph_format.space_before = Pt(4)
                p_proj.paragraph_format.space_after = Pt(2)
                
                run_title = p_proj.add_run(f"{title}")
                run_title.font.name = font_name
                run_title.font.size = Pt(10.5)
                run_title.font.bold = True
                run_title.font.color.rgb = secondary_color

                if tech:
                    run_tech = p_proj.add_run(f" [{tech}]")
                    run_tech.font.name = font_name
                    run_tech.font.size = Pt(9.5)
                    run_tech.font.italic = True
                    run_tech.font.color.rgb = primary_color

                bullets = proj.get("Bullet_Points") or proj.get("Description") or []
                if isinstance(bullets, list):
                    for b in bullets:
                        bp = doc.add_paragraph(style='List Bullet')
                        bp.paragraph_format.space_after = Pt(2)
                        run_b = bp.add_run(str(b))
                        run_b.font.name = font_name
                        run_b.font.size = Pt(9.5)
                        run_b.font.color.rgb = secondary_color
                elif bullets:
                    bp = doc.add_paragraph(style='List Bullet')
                    bp.paragraph_format.space_after = Pt(2)
                    run_b = bp.add_run(str(bullets))
                    run_b.font.name = font_name
                    run_b.font.size = Pt(9.5)
            else:
                p_proj = doc.add_paragraph(style='List Bullet')
                run_p = p_proj.add_run(str(proj))
                run_p.font.name = font_name
                run_p.font.size = Pt(9.5)

    # 7. Certifications
    certs = data.get("Certifications", [])
    if certs:
        add_section_header("Certifications & Licenses")
        for cert in certs:
            p_cert = doc.add_paragraph(style='List Bullet')
            p_cert.paragraph_format.space_after = Pt(2)
            if isinstance(cert, dict):
                cert_title = cert.get("Certificate", "Certification")
                inst = cert.get("Institution", "")
                text = f"{cert_title} — {inst}" if inst else cert_title
            else:
                text = str(cert)
            run_c = p_cert.add_run(text)
            run_c.font.name = font_name
            run_c.font.size = Pt(9.5)

    # 8. Education
    edu_list = data.get("Education", [])
    if edu_list:
        add_section_header("Education Credentials")
        for edu in edu_list:
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_before = Pt(3)
            p_edu.paragraph_format.space_after = Pt(2)
            if isinstance(edu, dict):
                deg = edu.get("Degree", "Degree")
                inst = edu.get("Institution", "")
                session = edu.get("Session", "")
                score = edu.get("Score", "")
                text = f"{deg} — {inst}"
                if session:
                    text += f" ({session})"
                if score:
                    text += f" | Score: {score}"
                run_ed = p_edu.add_run(text)
                run_ed.font.name = font_name
                run_ed.font.size = Pt(10)
                run_ed.font.bold = True
                run_ed.font.color.rgb = secondary_color

                desc = edu.get("Description") or edu.get("Highlights")
                if desc:
                    bp = doc.add_paragraph(style='List Bullet')
                    run_d = bp.add_run(str(desc))
                    run_d.font.name = font_name
                    run_d.font.size = Pt(9.5)
            else:
                run_ed = p_edu.add_run(str(edu))
                run_ed.font.name = font_name
                run_ed.font.size = Pt(9.5)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
